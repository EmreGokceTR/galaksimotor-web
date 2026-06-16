"""
GALAKSİ MOTOR - GÜNCEL BİLGİLER.docx üretici.
Çalıştırma: python scripts/generate_guncel_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT_PATH = os.path.join(os.path.expanduser("~"), "Desktop", "GALAKSİ MOTOR - GÜNCEL BİLGİLER.docx")

doc = Document()

# Tema
styles = doc.styles
style = styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xB8, 0x8A, 0x00)


def h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def kv(key, val):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{key}: ")
    r1.bold = True
    p.add_run(val)


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def hr():
    doc.add_paragraph("─" * 60)


# ── Başlık
h1("GALAKSİ MOTOR — GÜNCEL BİLGİLER")
p = doc.add_paragraph()
r = p.add_run("Son güncelleme: 16 Haziran 2026")
r.italic = True
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
hr()

# ── Site
h2("1. Site & Alan Adı")
kv("Alan adı (canlı)", "https://galaksimotor.com")
kv("WWW yönlendirmesi", "www.galaksimotor.com → galaksimotor.com (308)")
kv("Vercel projesi", "galaksimotor-web (team: galaksi-motors-projects)")
kv("Hosting", "Vercel (Edge + ISR + Cron)")
kv("Veritabanı (canlı)", "MySQL — Aiven Cloud (galaksimotor-galaksimotor.h.aivencloud.com)")
kv("Dosya/medya", "Vercel Blob (vercel_blob_rw_… tokenıyla)")
kv("E-posta", "Zoho Mail — smtp.zoho.eu:465 (SSL)")
kv("Yönetici e-posta", "info@galaksimotor.com")

# ── İyzico
h2("2. İyzico Ödeme Entegrasyonu  (CANLI — 16.06.2026)")
kv("Mod", "CANLI (PRODUCTION)")
kv("Base URL", "https://api.iyzipay.com")
kv("API Anahtarı", "gQgU3fNZTsLkjjqrpb6aZUEaH3hboBdO")
kv("Güvenlik Anahtarı", "jXFj09oYL3ijM4hmVI9UXiQy93IM131S")
kv("Entegrasyon Tipi", "İyzico Checkout Form (Hosted)")
kv("Kullanılan SDK", "iyzipay (npm) — src/lib/iyzico.ts")
kv("Korumalı Alışveriş", "Aynı API anahtarlarıyla çalışır; İyzico panelinden 'Korumalı Alışveriş' ürünü aktif edilmişse otomatik kullanılır. Ek bir 'işyeri numarası' alanı yoktur — merchant hesap düzeyinde tanımlıdır.")
kv("Vercel env (Production)", "IYZICO_API_KEY, IYZICO_SECRET_KEY, IYZICO_BASE_URL — hepsi ayarlandı")
kv("Geri ödeme & iptal", "Aynı SDK üzerinden cancel/refund destekli (admin paneli)")
kv("Test kart (sandbox modunda)", "5528790000000008 / 12/30 / 123 (Halkbank)")

h3("İyzico — Önemli Notlar")
bullet("Canlı mod aktif: gerçek kart çekimleri başlayacak. İlk siparişlerden sonra İyzico panelinde 'İşlemlerim' ekranını kontrol et.")
bullet("3D Secure varsayılan olarak aktif.")
bullet("İade ve iptal işlemleri admin panelinde /admin/siparisler üzerinden yapılır.")
bullet("Webhook/callback URL: https://galaksimotor.com/api/payment/iyzico/callback")

# ── E-posta Akışı
h2("3. Otomatik E-posta Akışı (Hepsi Aktif)")
bullet("Yeni üye kaydı → Hoşgeldin maili")
bullet("Şifre sıfırlama → Marka temalı sıfırlama linki (60 dk geçerli)")
bullet("Sipariş oluşturuldu → Müşteriye 'Sipariş alındı' bildirimi")
bullet("Ödeme başarılı → Müşteriye sipariş onayı + PDF FATURA EKLİ")
bullet("Ödeme başarılı → Admin'e 'Yeni Sipariş' alarmı")
bullet("Sipariş durumu değişti → Müşteriye durum bildirimi (HAZIRLANIYOR / KARGO / TESLİM)")
bullet("Randevu oluşturuldu → Müşteri + admin bildirimi")
bullet("Randevu durumu değişti → Müşteriye durum bildirimi")
bullet("İletişim formu → Müşteriye otomatik onay + admin'e içerik")
bullet("Düşük stok → Admin'e uyarı maili")

# ── E-ticaret
h2("4. E-ticaret Altyapısı")
bullet("Sepet (LocalStorage kalıcı) + drawer")
bullet("Üyelik: Google OAuth + e-posta/şifre")
bullet("Kupon kodu sistemi (yüzde / tutar / koşul)")
bullet("Stok takibi + otomatik düşürme")
bullet("KDV %20 ayrımlı PDF fatura (pdfkit ile)")
bullet("Vercel Blob'da arşivlenir, müşteri mailine ek olarak gönderilir")
bullet("Kargo: 49.90 ₺ sabit, 1-3 iş günü (Yurtiçi/Aras)")
bullet("İade hakkı: 14 gün (Mesafeli Satış Sözleşmesi'ne uygun)")
bullet("Online randevu sistemi (servis için saat seçimi, çakışma kontrolü)")

# ── SEO
h2("5. SEO & Pazarlama")
kv("Google Search Console", "Tek mülk (galaksimotor.com) — www mülkü kaldırıldı")
kv("Sitemap", "https://galaksimotor.com/sitemap.xml (otomatik)")
kv("Robots", "https://galaksimotor.com/robots.txt")
kv("Google Analytics 4", "G-58R7QEQ2CE")
kv("IndexNow", "Aktif — Bing & Yandex'e otomatik haber verir")
kv("Schema.org", "Product, Breadcrumb, AggregateRating, Organization, WebSite (SearchAction), LocalBusiness, FAQ")
kv("Open Graph", "product:price, availability, condition, brand etiketleri")
kv("PWA Manifest", "Aktif — 'Ana ekrana ekle' desteği")

# ── Politika sayfaları
h2("6. Politika & Hukuki Sayfalar")
bullet("Gizlilik Politikası — /gizlilik-politikasi")
bullet("KVKK Aydınlatma Metni — /kvkk")
bullet("Mesafeli Satış Sözleşmesi — /mesafeli-satis-sozlesmesi")
bullet("İptal & İade Koşulları (14 gün) — /iptal-iade-kosullari")
bullet("Çerez Politikası + onay banner")
bullet("İletişim — /iletisim (form + harita)")
bullet("Kargo bilgisi — /kargo")
bullet("Sık Sorulan Sorular — /sss (DB editable, JSON-LD)")

# ── Yedekleme
h2("7. Yedekleme & Güvenlik")
kv("Otomatik DB yedek", "Vercel Cron — her gün 03:00 (Aiven full dump → Blob)")
kv("Manuel yedek", "Admin panel → JSON indir")
kv("Cron secret", "Vercel env (CRON_SECRET)")
kv("Rate limit", "Sliding window (in-memory) — register/appointment/forgot-password")
kv("Bcrypt", "12 round password hashing")
kv("CSRF & enumeration koruması", "Aktif (forgot-password identical response)")
kv("Aktivite logu (ActivityLog)", "Tüm admin işlemleri + e-postalar kaydedilir")

# ── Admin
h2("8. Admin Paneli Özellikleri")
bullet("Ürün CRUD + çoklu görsel + kategori yönetimi")
bullet("Sipariş yönetimi + durum değişimi + fatura indirme + iade/iptal")
bullet("Randevu yönetimi + durum değişimi")
bullet("Kupon CRUD")
bullet("Servis CRUD")
bullet("Yorum moderasyonu")
bullet("İletişim formu yanıtlama")
bullet("Kullanıcı listesi + arama")
bullet("Aktivite logu görüntüleme")
bullet("Yedekleme")

# ── Önemli env değişkenleri özet
h2("9. Önemli Ortam Değişkenleri (Vercel Production)")
data = [
    ("DATABASE_URL", "Aiven MySQL bağlantı stringi"),
    ("NEXTAUTH_URL", "https://galaksimotor.com"),
    ("NEXTAUTH_SECRET", "Vercel'de güvenli"),
    ("GOOGLE_CLIENT_ID / SECRET", "OAuth"),
    ("IYZICO_API_KEY", "gQgU3fNZTsLkjjqrpb6aZUEaH3hboBdO"),
    ("IYZICO_SECRET_KEY", "jXFj09oYL3ijM4hmVI9UXiQy93IM131S"),
    ("IYZICO_BASE_URL", "https://api.iyzipay.com (CANLI)"),
    ("SMTP_HOST/PORT/USER/PASS", "Zoho Mail (info@galaksimotor.com)"),
    ("ADMIN_EMAIL", "info@galaksimotor.com"),
    ("BLOB_READ_WRITE_TOKEN", "Vercel Blob"),
    ("CRON_SECRET", "Cron yetkilendirme"),
    ("NEXT_PUBLIC_GA_ID", "G-58R7QEQ2CE"),
    ("NEXT_PUBLIC_SITE_URL", "https://galaksimotor.com"),
    ("MAINTENANCE_MODE", "false"),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Değişken"
hdr[1].text = "Değer / Açıklama"
for k, v in data:
    row = table.add_row().cells
    row[0].text = k
    row[1].text = v

# ── Son durum
doc.add_paragraph()
h2("10. Şu Anki Durum (16.06.2026)")
p = doc.add_paragraph()
r = p.add_run("✅ Site canlıda, e-ticaret tam aktif. İyzico canlı modda. Test siparişiyle (küçük tutarla) doğrulama önerilir.")
r.bold = True
r.font.color.rgb = RGBColor(0x00, 0x66, 0x00)

doc.save(OUT_PATH)
print(f"OK: {OUT_PATH}")
